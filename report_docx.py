# -*- coding: utf-8 -*-
"""
report_docx.py
Generates Word (.docx) reports from an AnalysisResult (see analyzer.py),
mirroring report_pdf.py so both formats offer the same two report types:

  - "technical": full detail — every extracted field and every finding
    with its weight/severity. For security analysts / IT.
  - "executive": a short, jargon-free summary — verdict, recommendation,
    and the handful of findings that matter. For managers and
    non-technical stakeholders.

Uses python-docx (pip install python-docx).

Part of E-MailX-Ray.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from report_common import build_summary, SEVERITY_LABELS

RISK_COLORS_RGB = {
    "High": RGBColor(0xC0, 0x39, 0x2B),
    "Medium": RGBColor(0xE6, 0x7E, 0x22),
    "Low": RGBColor(0x27, 0xAE, 0x60),
}

SEVERITY_COLORS_RGB = {
    "high": RGBColor(0xC0, 0x39, 0x2B),
    "medium": RGBColor(0xE6, 0x7E, 0x22),
    "low": RGBColor(0x9A, 0x7C, 0x0C),
    "info": RGBColor(0x2C, 0x3E, 0x50),
    "ai": RGBColor(0x8E, 0x44, 0xAD),
}

GREY = RGBColor(0x78, 0x78, 0x78)
DARK = RGBColor(0x14, 0x14, 0x14)
BODY = RGBColor(0x46, 0x46, 0x46)


# ----------------------------------------------------------------------
# Low-level helpers
# ----------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = DARK
    p.space_after = Pt(2)
    return p


def _subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY
    p.space_after = Pt(10)
    return p


def _risk_banner(doc, summary):
    color = RISK_COLORS_RGB.get(summary["risk_level"], RGBColor(0x7F, 0x7F, 0x7F))
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, _rgb_to_hex(color))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        f"  Risk: {summary['risk_level']}    |    Score: {summary['score']}    |    "
        f"{summary['total_findings']} finding(s)  "
    )
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for c in table.columns:
        c.width = Cm(16)
    doc.add_paragraph().space_after = Pt(2)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.border = None
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK
    # bottom border for a rule-like divider
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D2D2D2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def _body_paragraph(doc, text, size=10.5, color=BODY, italic=False, bold=False, indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    run.bold = bold
    return p


def _callout_box(doc, label, body, color):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    for c in table.columns:
        c.width = Cm(16)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:color"), _rgb_to_hex(color))
        borders.append(el)
    tc_pr.append(borders)

    p1 = cell.paragraphs[0]
    r1 = p1.add_run(label)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = color

    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BODY
    doc.add_paragraph().space_after = Pt(2)


def _finding_block(doc, f, indent=0.0):
    sev_color = SEVERITY_COLORS_RGB.get(f.severity, BODY)
    label = SEVERITY_LABELS.get(f.severity, "")
    suffix = f"   (+{f.weight} pts)" if f.weight else ""
    _body_paragraph(doc, f"[{label}] {f.rule}{suffix}", size=11, color=sev_color, bold=True, indent=indent)
    _body_paragraph(doc, f.detail, size=10, color=BODY, indent=indent + 0.4)


# ----------------------------------------------------------------------
# Executive report
# ----------------------------------------------------------------------

def _build_executive_docx(doc, result, summary):
    _title(doc, "Email Phishing Risk \u2014 Executive Summary")
    _subtitle(doc, f"Generated with {summary['app_name']} on {summary['generated_at']}")
    _risk_banner(doc, summary)

    _body_paragraph(doc, summary["verdict"], size=12.5, color=DARK)

    rec_color = RISK_COLORS_RGB.get(summary["risk_level"], BODY)
    _callout_box(doc, "Recommended action", summary["recommendation"], rec_color)

    _section_heading(doc, "Why")
    if summary["top_findings"]:
        for f in summary["top_findings"]:
            _finding_block(doc, f)
    else:
        _body_paragraph(doc, "No findings were raised by the analysis.")

    c = summary["counts"]
    _body_paragraph(
        doc,
        f"Full breakdown: {c['high']} high, {c['medium']} medium, {c['low']} low, "
        f"{c['ai']} AI-flagged, {c['info']} informational finding(s). "
        "See the technical report for complete details.",
        size=9, color=GREY, italic=True,
    )


# ----------------------------------------------------------------------
# Technical report
# ----------------------------------------------------------------------

def _build_technical_docx(doc, result, summary, raw_header=None):
    _title(doc, "Email Header Analysis Report \u2014 Technical Detail")
    _subtitle(doc, f"Generated with {summary['app_name']} on {summary['generated_at']}")
    _risk_banner(doc, summary)
    _body_paragraph(doc, summary["verdict"], size=10.5, color=GREY, italic=True)

    _section_heading(doc, "Extracted fields")
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for k, v in result.fields.items():
        row = table.add_row().cells
        row[0].width = Cm(4.5)
        row[1].width = Cm(11.5)
        r0 = row[0].paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = DARK
        r1 = row[1].paragraphs[0].add_run(str(v))
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = BODY
    doc.add_paragraph().space_after = Pt(4)

    _section_heading(doc, "Risk findings")
    for f in result.findings:
        _finding_block(doc, f)

    if raw_header:
        doc.add_page_break()
        _section_heading(doc, "Appendix: raw header analyzed")
        p = doc.add_paragraph()
        run = p.add_run(raw_header)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x3C, 0x3C, 0x3C)


# ----------------------------------------------------------------------
# Public entry point
# ----------------------------------------------------------------------

def generate_docx_report(result, output_path: str, app_name: str = "E-MailX-Ray",
                          report_type: str = "technical", raw_header: str = None):
    """
    result: instance of analyzer.AnalysisResult
    output_path: path to save the .docx to
    report_type: "technical" (full detail) or "executive" (short summary)
    raw_header: optional raw header text, included as an appendix in the
                technical report only
    """
    summary = build_summary(result, app_name=app_name)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    if report_type == "executive":
        _build_executive_docx(doc, result, summary)
    else:
        _build_technical_docx(doc, result, summary, raw_header=raw_header)

    doc.save(output_path)
